import docx
from django.http import HttpResponse
from django.shortcuts import render
from django.views.generic import ListView, TemplateView, DetailView
from django.views.generic.edit import CreateView, UpdateView, DeleteView
from django.urls import reverse_lazy
from .models import Document
from pydocx import PyDocX
from .models import Document, DocumentConsent, Sample, DocumentNotification, DocumentAuthorsAward, DocumentSet, \
    ReviewProblemDocumentConsent, ReviewProblemSample, SignDocument
from .forms import SampleForm, DocumentConsentForm, DocumentNotificationForm, DocumentAuthorsAwardForm, DocumentSetForm, \
    ReviewProblemDocumentConsentForm, ReviewProblemSampleForm, SignDocumentForm, HeadSignDocumentForm
from django.views.generic.base import View
from django.http import FileResponse
from django.conf import settings
import os
from django.http import HttpResponse, HttpResponseBadRequest
from docx2pdf import convert
from docxtpl import DocxTemplate
from docx import Document
import tempfile


# Create your views here.

class DocumentSetListView(ListView):
    model = DocumentSet
    template_name = 'document_doc_pack/documents_sets_all_detail.html'


# Create your views here.


class DocumentAuthorsAwardListView(ListView):
    model = DocumentAuthorsAward
    template_name = 'document_authors_award/document_authors_award_all.html'


# Create your views here.
class DocumentNotificationListView(ListView):
    model = DocumentNotification
    template_name = 'document_notification/document_notification_all.html'


class DocumentConsentListView(ListView):
    model = DocumentConsent
    template_name = 'documents_consent_all.html'


# ----------------create view---------------- #
class DocumentSetsCreateView(CreateView):
    model = DocumentSet
    form_class = DocumentSetForm
    template_name = 'document_doc_pack/documents_sets_new.html'


class DocumentAuthorsAwardCreateView(CreateView):
    model = DocumentAuthorsAward
    form_class = DocumentAuthorsAwardForm
    template_name = 'document_authors_award/document_authors_award_new.html'


class DocumentNotificationCreateView(CreateView):
    model = DocumentNotification
    form_class = DocumentNotificationForm
    template_name = 'document_notification/document_notification_new.html'

    path_to_template = model.path_to_template


class DocumentConsentCreateView(CreateView):
    model = DocumentConsent
    form_class = DocumentConsentForm
    template_name = 'documents_consent_new.html'

    path_to_template = model.path_to_template


# ----------------detail view---------------- #

class DocumentSetsDetailView(DetailView):
    model = DocumentSet
    # form_class = DocumentSetForm
    template_name = 'document_doc_pack/document_set_detail.html'


class DocumentAuthorsAwardDetailView(DetailView):
    model = DocumentAuthorsAward
    template_name = 'document_authors_award/document_authors_award_detail.html'


class DocumentNotificationDetailView(DetailView):
    model = DocumentNotification
    template_name = 'document_notification/document_notification_detail.html'


class DocumentConsentDetailView(DetailView):
    model = DocumentConsent
    template_name = 'document_consent_detail.html'


# ----------------update view---------------- #


class DocumentSetsUpdateView(UpdateView):
    model = DocumentSet
    form_class = DocumentSetForm
    template_name = 'document_doc_pack/document_set_edit.html'


class DocumentAuthorsAwardUpdateView(UpdateView):
    model = DocumentAuthorsAward
    form_class = DocumentAuthorsAwardForm
    template_name = 'document_authors_award/document_authors_award_edit.html'


class DocumentNotificationUpdateView(UpdateView):
    model = DocumentNotification
    form_class = DocumentNotificationForm
    template_name = 'document_notification/document_notification_edit.html'


class DocumentConsentUpdateView(UpdateView):
    model = DocumentConsent
    form_class = DocumentConsentForm
    template_name = 'document_consent_edit.html'


# ----------------delete view---------------- #


class DocumentSetsDeleteView(DeleteView):
    model = DocumentSet
    template_name = 'document_doc_pack/document_set_delete.html'
    success_url = reverse_lazy('document_sets_all_detail')


class DocumentAuthorsAwardDeleteView(DeleteView):
    model = DocumentAuthorsAward
    template_name = 'document_authors_award/document_authors_award_delete.html'
    success_url = reverse_lazy('documents_authors_award_all')


class DocumentNotificationDeleteView(DeleteView):
    model = DocumentNotification
    template_name = 'document_notification/document_notification_delete.html'
    success_url = reverse_lazy('documents_notification_all')


class DocumentConsentDeleteView(DeleteView):
    model = DocumentConsent
    template_name = 'document_consent_delete.html'
    success_url = reverse_lazy('documents_consent_all')


class DocumentConsentDownloadDocx(View):
    def get(self, request, pk=1, *args, **kwargs):
        document_consent_program_name = DocumentConsent.objects.get(pk=pk).program_name
        document_consent_application_number = DocumentConsent.objects.get(pk=pk).application_number
        document_consent_applicant_name = DocumentConsent.objects.get(pk=pk).applicant_name
        document_consent_applicant_surname = DocumentConsent.objects.get(pk=pk).applicant_surname
        document_consent_applicant_patronomic = DocumentConsent.objects.get(pk=pk).applicant_patronomic
        document_consent_applicant_date_of_birth = DocumentConsent.objects.get(pk=pk).applicant_date_of_birth
        document_consent_address_index = DocumentConsent.objects.get(pk=pk).address_index
        document_consent_address_country = DocumentConsent.objects.get(pk=pk).address_country
        document_consent_address_city = DocumentConsent.objects.get(pk=pk).address_city
        document_consent_address_street = DocumentConsent.objects.get(pk=pk).address_street
        document_consent_address_building_number = DocumentConsent.objects.get(pk=pk).address_building_number
        document_consent_address_house_flat_number = DocumentConsent.objects.get(pk=pk).address_house_flat_number
        document_consent_applicant_phone_number = DocumentConsent.objects.get(pk=pk).applicant_phone_number
        document_consent_current_date = DocumentConsent.objects.get(pk=pk).current_date
        document_consent_passport_seria = DocumentConsent.objects.get(pk=pk).passport_seria
        document_consent_passport_number = DocumentConsent.objects.get(pk=pk).passport_number
        document_consent_passport_date_of_issue = DocumentConsent.objects.get(pk=pk).passport_date_of_issue
        document_consent_passport_place_giving = DocumentConsent.objects.get(pk=pk).passport_place_giving

        template_name_id = DocumentConsent.objects.get(pk=pk).template_name_id
        path_to_template = Sample.objects.get(pk=template_name_id).path_to_template.name

        file_name_from_path = os.path.basename(path_to_template)

        file_name_without_extension = str(os.path.splitext(file_name_from_path)[0])
        # file_name_without_extension = 'document_consent'

        # file_path = f'/Users/keito/Programming/Python/train/diploma/media/upload_sample/2024-05-07/{file_name_from_path}'
        # file_path = '/Users/keito/Programming/Python/train/diploma/media/upload_sample/2024-05-07/document_consent.docx'
        file_path = f'/Users/keito/Programming/Python/train/diploma/media/{path_to_template}'

        file_path_short = file_name_without_extension
        if file_path.endswith('.docx'):
            from datetime import datetime as dt
            doc = DocxTemplate(file_path)
            document2 = DocxTemplate(file_path)
            context = {
                'program_name': document_consent_program_name,
                'application_number': document_consent_application_number,
                'applicant_name': document_consent_applicant_name,
                'applicant_surname': document_consent_applicant_surname,
                'applicant_patronomic': document_consent_applicant_patronomic,
                'applicant_date_of_birth': dt.strftime(document_consent_applicant_date_of_birth, '%d.%m.%Y'),
                'birth_day': document_consent_applicant_date_of_birth.day,
                'birth_month': document_consent_applicant_date_of_birth.month if document_consent_applicant_date_of_birth.month >= 10 else str(
                    '0' + f'{document_consent_applicant_date_of_birth.month}'),
                'birth_year': document_consent_applicant_date_of_birth.year,
                'address_index': document_consent_address_index,
                'address_country': document_consent_address_country,
                'country_short': 'РФ',
                'address_city': document_consent_address_city,
                'address_street': document_consent_address_street,
                'address_building_number': document_consent_address_building_number,
                'address_house_flat_number': document_consent_address_house_flat_number,
                'applicant_phone_number': document_consent_applicant_phone_number,
                'current_date': dt.strftime(document_consent_current_date, '%d.%m.%Y'),
                'applicant_name_short': document_consent_applicant_name[0],
                'applicant_patronomic_short': document_consent_applicant_patronomic[0],
                'passport_seria': document_consent_passport_seria,
                'passport_number': document_consent_passport_number,
                'passport_date_of_issue': dt.strftime(document_consent_passport_date_of_issue, '%d.%m.%Y'),
                'passport_place_giving': document_consent_passport_place_giving,
            }
            # import aspose.words as aw
            # import random
            #
            # doc.render(context)
            # document2.render(context)
            # # file_name = f'/Users/keito/Downloads/{file_path_short}_downloaded.docx'
            # file_name = f'{file_path_short}_downloaded.docx'
            # file_name_2 = f'{file_path_short}_downloaded.docx'
            # # document = aw.Document(file_name)
            # document2 = aw.Document(file_name_2)
            #
            # doc.save(file_name)
            # document2.save(file_name_2)

            doc.render(context)
            # file_name = f'/Users/keito/Downloads/{file_path_short}_downloaded.docx'
            file_name = f'{file_path_short}_downloaded.docx'
            doc.save(file_name)

        response = FileResponse(open(file_name, 'rb'))
        response['Content-Disposition'] = f'attachment; filename="{file_name}"'
        return response


class DocumentConsentDownloadPDF(View):
    def get(self, request, pk=1, *args, **kwargs):
        document_consent_program_name = DocumentConsent.objects.get(pk=pk).program_name
        document_consent_application_number = DocumentConsent.objects.get(pk=pk).application_number
        document_consent_applicant_name = DocumentConsent.objects.get(pk=pk).applicant_name
        document_consent_applicant_surname = DocumentConsent.objects.get(pk=pk).applicant_surname
        document_consent_applicant_patronomic = DocumentConsent.objects.get(pk=pk).applicant_patronomic
        document_consent_applicant_date_of_birth = DocumentConsent.objects.get(pk=pk).applicant_date_of_birth
        document_consent_address_index = DocumentConsent.objects.get(pk=pk).address_index
        document_consent_address_country = DocumentConsent.objects.get(pk=pk).address_country
        document_consent_address_city = DocumentConsent.objects.get(pk=pk).address_city
        document_consent_address_street = DocumentConsent.objects.get(pk=pk).address_street
        document_consent_address_building_number = DocumentConsent.objects.get(pk=pk).address_building_number
        document_consent_address_house_flat_number = DocumentConsent.objects.get(pk=pk).address_house_flat_number
        document_consent_applicant_phone_number = DocumentConsent.objects.get(pk=pk).applicant_phone_number
        document_consent_current_date = DocumentConsent.objects.get(pk=pk).current_date
        document_consent_passport_seria = DocumentConsent.objects.get(pk=pk).passport_seria
        document_consent_passport_number = DocumentConsent.objects.get(pk=pk).passport_number
        document_consent_passport_date_of_issue = DocumentConsent.objects.get(pk=pk).passport_date_of_issue
        document_consent_passport_place_giving = DocumentConsent.objects.get(pk=pk).passport_place_giving

        template_name_id = DocumentConsent.objects.get(pk=pk).template_name_id
        path_to_template = Sample.objects.get(pk=template_name_id).path_to_template.name

        file_name_from_path = os.path.basename(path_to_template)

        file_name_without_extension = str(os.path.splitext(file_name_from_path)[0])
        # file_name_without_extension = 'document_consent'

        # file_path = f'/Users/keito/Programming/Python/train/diploma/media/upload_sample/2024-05-07/{file_name_from_path}'
        # file_path = '/Users/keito/Programming/Python/train/diploma/media/upload_sample/2024-05-07/document_consent.docx'
        file_path = f'/Users/keito/Programming/Python/train/diploma/media/{path_to_template}'

        file_path_short = file_name_without_extension
        if file_path.endswith('.docx'):
            from datetime import datetime as dt
            doc = DocxTemplate(file_path)
            context = {
                'program_name': document_consent_program_name,
                'application_number': document_consent_application_number,
                'applicant_name': document_consent_applicant_name,
                'applicant_surname': document_consent_applicant_surname,
                'applicant_patronomic': document_consent_applicant_patronomic,
                'applicant_date_of_birth': dt.strftime(document_consent_applicant_date_of_birth, '%d.%m.%Y'),
                'birth_day': document_consent_applicant_date_of_birth.day,
                'birth_month': document_consent_applicant_date_of_birth.month if document_consent_applicant_date_of_birth.month >= 10 else str(
                    '0' + f'{document_consent_applicant_date_of_birth.month}'),
                'birth_year': document_consent_applicant_date_of_birth.year,
                'address_index': document_consent_address_index,
                'address_country': document_consent_address_country,
                'country_short': 'РФ',
                'address_city': document_consent_address_city,
                'address_street': document_consent_address_street,
                'address_building_number': document_consent_address_building_number,
                'address_house_flat_number': document_consent_address_house_flat_number,
                'applicant_phone_number': document_consent_applicant_phone_number,
                'current_date': dt.strftime(document_consent_current_date, '%d.%m.%Y'),
                'applicant_name_short': document_consent_applicant_name[0],
                'applicant_patronomic_short': document_consent_applicant_patronomic[0],
                'passport_seria': document_consent_passport_seria,
                'passport_number': document_consent_passport_number,
                'passport_date_of_issue': dt.strftime(document_consent_passport_date_of_issue, '%d.%m.%Y'),
                'passport_place_giving': document_consent_passport_place_giving,
            }

            doc.render(context)
            file_name = f'{file_path_short}_downloaded.docx'
            import aspose.words as aw
            file_name_pdf = f'{file_path_short}_downloaded.pdf'
            document = aw.Document(file_name)
            # doc.save(file_name)
            document.save(file_name_pdf)

        response = FileResponse(open(file_name_pdf, 'rb'))
        response['Content-Disposition'] = f'attachment; filename="{file_name_pdf}"'
        return response


# TODO: https://stackoverflow.com/questions/67385436/how-to-save-a-python-docxtemplate-as-pdf-quickly:
class DocumentNotificationDownloadDocx(View):
    def get(self, request, pk=1, *args, **kwargs):
        document_notification_program_name = DocumentNotification.objects.get(pk=pk).program_name
        document_notification_program_destiny = DocumentNotification.objects.get(pk=pk).program_destiny
        document_notification_head_name = DocumentNotification.objects.get(pk=pk).head_name
        document_notification_head_surname = DocumentNotification.objects.get(pk=pk).head_surname
        document_notification_head_patronomic = DocumentNotification.objects.get(pk=pk).head_patronomic
        document_notification_applicant1_name = DocumentNotification.objects.get(pk=pk).applicant1_name
        document_notification_applicant1_surname = DocumentNotification.objects.get(pk=pk).applicant1_surname
        document_notification_applicant1_patronomic = DocumentNotification.objects.get(pk=pk).applicant1_patronomic
        document_notification_applicant2_name = DocumentNotification.objects.get(pk=pk).applicant2_name
        document_notification_applicant2_surname = DocumentNotification.objects.get(pk=pk).applicant2_surname
        document_notification_applicant2_patronomic = DocumentNotification.objects.get(pk=pk).applicant2_patronomic
        document_notification_applicant1_uni_position = DocumentNotification.objects.get(pk=pk).applicant1_uni_position
        document_notification_applicant2_uni_position = DocumentNotification.objects.get(pk=pk).applicant2_uni_position
        document_notification_uni_department = DocumentNotification.objects.get(pk=pk).uni_department
        document_notification_applicant1_index = DocumentNotification.objects.get(pk=pk).applicant1_index
        document_notification_applicant1_country = DocumentNotification.objects.get(pk=pk).applicant1_country
        document_notification_applicant1_city = DocumentNotification.objects.get(pk=pk).applicant1_city
        document_notification_applicant1_street = DocumentNotification.objects.get(pk=pk).applicant1_street
        document_notification_applicant1_build_number = DocumentNotification.objects.get(pk=pk).applicant1_build_number
        document_notification_applicant1_flat_number = DocumentNotification.objects.get(pk=pk).applicant1_flat_number
        document_notification_applicant1_snils = DocumentNotification.objects.get(pk=pk).applicant1_snils
        document_notification_applicant2_index = DocumentNotification.objects.get(pk=pk).applicant2_index
        document_notification_applicant2_country = DocumentNotification.objects.get(pk=pk).applicant2_country
        document_notification_applicant2_city = DocumentNotification.objects.get(pk=pk).applicant2_city
        document_notification_applicant2_street = DocumentNotification.objects.get(pk=pk).applicant2_street
        document_notification_applicant2_build_number = DocumentNotification.objects.get(pk=pk).applicant2_build_number
        document_notification_applicant2_flat_number = DocumentNotification.objects.get(pk=pk).applicant2_flat_number
        document_notification_applicant2_snils = DocumentNotification.objects.get(pk=pk).applicant2_snils
        document_notification_applicant1_percent_contribution = DocumentNotification.objects.get(
            pk=pk).applicant1_percent_contribution
        document_notification_applicant2_percent_contribution = DocumentNotification.objects.get(
            pk=pk).applicant2_percent_contribution
        document_notification_applicant1_phone_number = DocumentNotification.objects.get(pk=pk).applicant1_phone_number
        document_notification_applicant2_phone_number = DocumentNotification.objects.get(pk=pk).applicant2_phone_number
        document_notification_applicant1_email = DocumentNotification.objects.get(pk=pk).applicant1_email
        document_notification_applicant2_email = DocumentNotification.objects.get(pk=pk).applicant2_email
        document_notification_program_usage = DocumentNotification.objects.get(pk=pk).program_usage
        document_notification_fee = DocumentNotification.objects.get(pk=pk).fee

        template_name_id = DocumentNotification.objects.get(pk=pk).template_name_id
        path_to_template = Sample.objects.get(pk=template_name_id).path_to_template.name

        file_name_from_path = os.path.basename(path_to_template)

        file_name_without_extension = str(os.path.splitext(file_name_from_path)[0])
        # file_name_without_extension = 'document_consent'

        # file_path = f'/Users/keito/Programming/Python/train/diploma/media/upload_sample/2024-05-07/{file_name_from_path}'
        # file_path = '/Users/keito/Downloads/document_notif.docx'
        file_path = f'/Users/keito/Programming/Python/train/diploma/media/{path_to_template}'

        file_path_short = file_name_without_extension
        if file_path.endswith('.docx'):
            from datetime import datetime as dt
            doc = DocxTemplate(file_path)
            context = {
                'program_name': document_notification_program_name,
                'program_destiny': document_notification_program_destiny,

                # # руководитель
                'head_name': document_notification_head_name,
                'head_surname': document_notification_head_surname,
                'head_patronomic': document_notification_head_patronomic,

                'applicant1_name': document_notification_applicant1_name,
                'applicant1_name_short': document_notification_applicant1_name[0],
                'applicant1_surname': document_notification_applicant1_surname,
                'applicant1_patronomic': document_notification_applicant1_patronomic,
                'applicant1_patronomic_short': document_notification_applicant1_patronomic[0],

                #
                'applicant2_name': document_notification_applicant2_name,
                'applicant2_name_short': document_notification_applicant2_name[0],
                'applicant2_surname': document_notification_applicant2_surname,
                'applicant2_patronomic': document_notification_applicant2_patronomic,
                'applicant2_patronomic_short': document_notification_applicant2_patronomic[0],

                #
                'applicant1_uni_position': document_notification_applicant1_uni_position,
                'applicant2_uni_position': document_notification_applicant2_uni_position,
                #
                'uni_department': document_notification_uni_department,
                #
                'applicant1_index': document_notification_applicant1_index,
                'applicant1_country': document_notification_applicant1_country,
                'applicant1_city': document_notification_applicant1_city,
                'applicant1_street': document_notification_applicant1_street,
                'applicant1_build_number': document_notification_applicant1_build_number,
                'applicant1_flat_number': document_notification_applicant1_flat_number,
                'applicant1_snils': document_notification_applicant1_snils,
                #
                'applicant2_index': document_notification_applicant2_index,
                'applicant2_country': document_notification_applicant2_country,
                'applicant2_city': document_notification_applicant2_city,
                'applicant2_street': document_notification_applicant2_street,
                'applicant2_build_number': document_notification_applicant2_build_number,
                'applicant2_flat_number': document_notification_applicant2_flat_number,
                'applicant2_snils': document_notification_applicant2_snils,
                #
                'applicant1_percent_contribution': document_notification_applicant1_percent_contribution,
                'applicant2_percent_contribution': document_notification_applicant2_percent_contribution,
                #
                'applicant1_phone_number': document_notification_applicant1_phone_number,
                'applicant2_phone_number': document_notification_applicant2_phone_number,
                #
                'applicant1_email': document_notification_applicant1_email,
                'applicant2_email': document_notification_applicant2_email,
                'program_usage': document_notification_program_usage,
                'fee': document_notification_fee
            }

            doc.render(context)
            # file_name = f'/Users/keito/Downloads/{file_path_short}_downloaded.docx'
            file_name = f'{file_path_short}_downloaded.docx'
            doc.save(file_name)

        response = FileResponse(open(file_name, 'rb'))
        response['Content-Disposition'] = f'attachment; filename="{file_name}"'
        return response


class DocumentAuthorsAwardDownloadDocx(View):
    def get(self, request, pk=1, *args, **kwargs):
        document_authors_award_program_name = DocumentAuthorsAward.objects.get(pk=pk).program_name

        document_authors_award_applicant1_name = DocumentAuthorsAward.objects.get(pk=pk).applicant1_name
        document_authors_award_applicant1_surname = DocumentAuthorsAward.objects.get(pk=pk).applicant1_surname
        document_authors_award_applicant1_patronomic = DocumentAuthorsAward.objects.get(pk=pk).applicant1_patronomic
        document_authors_award_applicant1_date_birth = DocumentAuthorsAward.objects.get(pk=pk).applicant1_date_birth

        document_authors_award_applicant1_passport_seria = DocumentAuthorsAward.objects.get(
            pk=pk).applicant1_passport_seria
        document_authors_award_applicant1_passport_number = DocumentAuthorsAward.objects.get(
            pk=pk).applicant1_passport_number
        document_authors_award_applicant1_passport_date_of_issue = DocumentAuthorsAward.objects.get(
            pk=pk).applicant1_passport_date_of_issue
        document_authors_award_applicant1_passport_place_giving = DocumentAuthorsAward.objects.get(
            pk=pk).applicant1_passport_place_giving

        document_authors_award_applicant1_inn = DocumentAuthorsAward.objects.get(pk=pk).applicant1_inn
        document_authors_award_applicant1_snils = DocumentAuthorsAward.objects.get(pk=pk).applicant1_snils
        document_authors_award_applicant1_bank_info = DocumentAuthorsAward.objects.get(pk=pk).applicant1_bank_info
        document_authors_award_applicant1_account_number = DocumentAuthorsAward.objects.get(
            pk=pk).applicant1_account_number
        document_authors_award_applicant1_bik = DocumentAuthorsAward.objects.get(pk=pk).applicant1_bik
        document_authors_award_applicant1_bank_inn = DocumentAuthorsAward.objects.get(pk=pk).applicant1_bank_inn
        document_authors_award_applicant1_bank_kpp = DocumentAuthorsAward.objects.get(pk=pk).applicant1_bank_kpp
        document_authors_award_applicant1_corr_account = DocumentAuthorsAward.objects.get(pk=pk).applicant1_corr_account

        document_authors_award_applicant2_name = DocumentAuthorsAward.objects.get(pk=pk).applicant2_name
        document_authors_award_applicant2_surname = DocumentAuthorsAward.objects.get(pk=pk).applicant2_surname
        document_authors_award_applicant2_patronomic = DocumentAuthorsAward.objects.get(pk=pk).applicant2_patronomic
        document_authors_award_applicant2_date_birth = DocumentAuthorsAward.objects.get(pk=pk).applicant2_date_birth

        document_authors_award_applicant2_passport_seria = DocumentAuthorsAward.objects.get(
            pk=pk).applicant2_passport_seria
        document_authors_award_applicant2_passport_number = DocumentAuthorsAward.objects.get(
            pk=pk).applicant2_passport_number
        document_authors_award_applicant2_passport_date_of_issue = DocumentAuthorsAward.objects.get(
            pk=pk).applicant2_passport_date_of_issue
        document_authors_award_applicant2_passport_place_giving = DocumentAuthorsAward.objects.get(
            pk=pk).applicant2_passport_place_giving

        document_authors_award_applicant2_inn = DocumentAuthorsAward.objects.get(pk=pk).applicant2_inn
        document_authors_award_applicant2_snils = DocumentAuthorsAward.objects.get(pk=pk).applicant2_snils
        document_authors_award_applicant2_bank_info = DocumentAuthorsAward.objects.get(pk=pk).applicant2_bank_info
        document_authors_award_applicant2_account_number = DocumentAuthorsAward.objects.get(
            pk=pk).applicant2_account_number
        document_authors_award_applicant2_bik = DocumentAuthorsAward.objects.get(pk=pk).applicant2_bik
        document_authors_award_applicant2_bank_inn = DocumentAuthorsAward.objects.get(pk=pk).applicant2_bank_inn
        document_authors_award_applicant2_bank_kpp = DocumentAuthorsAward.objects.get(pk=pk).applicant2_bank_kpp
        document_authors_award_applicant2_corr_account = DocumentAuthorsAward.objects.get(pk=pk).applicant2_corr_account

        document_authors_award_authors_award = DocumentAuthorsAward.objects.get(pk=pk).authors_award
        document_authors_award_authors_award_half = document_authors_award_authors_award // 2

        template_name_id = DocumentAuthorsAward.objects.get(pk=pk).template_name_id
        path_to_template = Sample.objects.get(pk=template_name_id).path_to_template.name

        file_name_from_path = os.path.basename(path_to_template)

        file_name_without_extension = str(os.path.splitext(file_name_from_path)[0])
        # file_name_without_extension = 'document_consent'

        # file_path = f'/Users/keito/Programming/Python/train/diploma/media/upload_sample/2024-05-07/{file_name_from_path}'
        # file_path = '/Users/keito/Downloads/document_notif.docx'
        file_path = f'/Users/keito/Programming/Python/train/diploma/media/{path_to_template}'

        file_path_short = file_name_without_extension
        if file_path.endswith('.docx'):
            from datetime import datetime as dt
            from numtostr_rus import convert
            doc = DocxTemplate(file_path)
            context = {
                'program_name': document_authors_award_program_name,

                'applicant1_name': document_authors_award_applicant1_name,
                'applicant1_surname': document_authors_award_applicant1_surname,
                'applicant1_patronomic': document_authors_award_applicant1_patronomic,
                'applicant1_date_birth': dt.strftime(document_authors_award_applicant1_date_birth, '%d.%m.%Y'),

                'applicant1_index': '413111',
                'applicant1_country': 'Российская Федерация',
                'applicant1_city': 'Энгельс',
                'applicant1_street': 'пр-т Фридриха Энгельса',
                'applicant1_building_number': 11,
                'applicant1_flat_number': 69,

                'applicant1_passport_seria': document_authors_award_applicant1_passport_seria,
                'applicant1_passport_number': document_authors_award_applicant1_passport_number,
                'applicant1_passport_date_of_issue': dt.strftime(
                    document_authors_award_applicant1_passport_date_of_issue, '%d.%m.%Y'),
                'applicant1_passport_place_giving': document_authors_award_applicant1_passport_place_giving,

                'applicant1_inn': document_authors_award_applicant1_inn,
                'applicant1_snils': document_authors_award_applicant1_snils,
                'applicant1_bank_info': document_authors_award_applicant1_bank_info,
                'applicant1_account_number': document_authors_award_applicant1_account_number,
                'applicant1_bik': document_authors_award_applicant1_bik,
                'applicant1_bank_inn': document_authors_award_applicant1_bank_inn,
                'applicant1_bank_kpp': document_authors_award_applicant1_bank_kpp,
                'applicant1_corr_account': document_authors_award_applicant1_corr_account,

                # 2
                'applicant2_name': document_authors_award_applicant2_name,
                'applicant2_surname': document_authors_award_applicant2_surname,
                'applicant2_patronomic': document_authors_award_applicant2_patronomic,
                'applicant2_date_birth': dt.strftime(document_authors_award_applicant2_date_birth, '%d.%m.%Y'),

                'applicant2_index': '999111',
                'applicant2_country': 'Российская Федерация',
                'applicant2_city': 'Саратов',
                'applicant2_street': 'Кирова',
                'applicant2_building_number': 1,
                'applicant2_flat_number': 629,

                'applicant2_passport_seria': document_authors_award_applicant2_passport_seria,
                'applicant2_passport_number': document_authors_award_applicant2_passport_number,
                'applicant2_passport_date_of_issue': dt.strftime(
                    document_authors_award_applicant2_passport_date_of_issue, '%d.%m.%Y'),
                'applicant2_passport_place_giving': document_authors_award_applicant2_passport_place_giving,

                'applicant2_inn': document_authors_award_applicant2_inn,
                'applicant2_snils': document_authors_award_applicant2_snils,
                'applicant2_bank_info': document_authors_award_applicant2_bank_info,
                'applicant2_account_number': document_authors_award_applicant2_account_number,
                'applicant2_bik': document_authors_award_applicant2_bik,
                'applicant2_bank_inn': document_authors_award_applicant2_bank_inn,
                'applicant2_bank_kpp': document_authors_award_applicant2_bank_kpp,
                'applicant2_corr_account': document_authors_award_applicant2_corr_account,

                'applicant1_name_short': document_authors_award_applicant1_name[0],
                'applicant1_patronomic_short': document_authors_award_applicant1_patronomic[0],
                'applicant2_name_short': document_authors_award_applicant2_name[0],
                'applicant2_patronomic_short': document_authors_award_applicant2_patronomic[0],

                'applicant1_authors_award': document_authors_award_authors_award_half,
                'applicant2_authors_award': document_authors_award_authors_award_half,

                'authors_award_text': convert(document_authors_award_authors_award),
                'applicant1_authors_award_text': convert(document_authors_award_authors_award_half),
                'applicant2_authors_award_text': convert(document_authors_award_authors_award_half),
            }

            doc.render(context)
            # file_name = f'/Users/keito/Downloads/{file_path_short}_downloaded.docx'
            file_name = f'{file_path_short}_downloaded.docx'
            doc.save(file_name)

        response = FileResponse(open(file_name, 'rb'))
        response['Content-Disposition'] = f'attachment; filename="{file_name}"'
        return response


class SampleListView(ListView):
    model = Sample
    template_name = 'samples_list.html'


class SampleDetailView(DetailView):
    model = Sample
    template_name = 'sample_detail.html'


class FileDownloadDocx(View):
    def get(self, request, pk=1, *args, **kwargs):
        sample = Sample.objects.get(pk=pk).path_to_template
        file_path = os.path.join(settings.MEDIA_ROOT, f'{sample}')  # Specify the path to your file
        if file_path.endswith('.docx'):
            file_name = f'{sample}_downloaded.docx'

        response = FileResponse(open(file_path, 'rb'))
        response['Content-Disposition'] = f'attachment; filename="{file_name}"'
        return response


class FileDownloadPdf(View):
    def get(self, request, pk=1, *args, **kwargs):
        sample = Sample.objects.get(pk=pk).path_to_template
        file_path = os.path.join(settings.MEDIA_ROOT, f'{sample}')  # Specify the path to your file
        if file_path.endswith('.docx'):
            file_name_1 = 'downloaded_file.docx'
            file_name = convert(file_path, file_name_1)
        elif file_path.endswith('.pdf'):
            file_name = 'downloaded_file.pdf'
        # else:
        #     return HttpResponseBadRequest('Unsupported file format. Only DOCX and PDF files are allowed.'
        #
        response = FileResponse(open(file_path, 'rb'))
        response['Content-Disposition'] = f'attachment; filename="{file_name}"'
        return response


class DocumentSetsDownloadDocx(View):
    def get(self, request, pk=1, *args, **kwargs):
        document_consent_program_name = DocumentSet.objects.get(pk=pk).document_consent.program_name
        # DocumentConsent.objects.get(pk=pk).program_name
        document_consent_application_number = DocumentSet.objects.get(pk=pk).document_consent.application_number
        document_consent_applicant_name = DocumentSet.objects.get(pk=pk).document_consent.applicant_name
        document_consent_applicant_surname = DocumentSet.objects.get(pk=pk).document_consent.applicant_surname
        document_consent_applicant_patronomic = DocumentSet.objects.get(pk=pk).document_consent.applicant_patronomic
        document_consent_applicant_date_of_birth = DocumentSet.objects.get(
            pk=pk).document_consent.applicant_date_of_birth
        document_consent_address_index = DocumentSet.objects.get(pk=pk).document_consent.address_index
        document_consent_address_country = DocumentSet.objects.get(pk=pk).document_consent.address_country
        document_consent_address_city = DocumentSet.objects.get(pk=pk).document_consent.address_city
        document_consent_address_street = DocumentSet.objects.get(pk=pk).document_consent.address_street
        document_consent_address_building_number = DocumentSet.objects.get(
            pk=pk).document_consent.address_building_number
        document_consent_address_house_flat_number = DocumentSet.objects.get(
            pk=pk).document_consent.address_house_flat_number
        document_consent_applicant_phone_number = DocumentSet.objects.get(pk=pk).document_consent.applicant_phone_number
        document_consent_current_date = DocumentSet.objects.get(pk=pk).document_consent.current_date
        document_consent_passport_seria = DocumentSet.objects.get(pk=pk).document_consent.passport_seria
        document_consent_passport_number = DocumentSet.objects.get(pk=pk).document_consent.passport_number
        document_consent_passport_date_of_issue = DocumentSet.objects.get(pk=pk).document_consent.passport_date_of_issue
        document_consent_passport_place_giving = DocumentSet.objects.get(pk=pk).document_consent.passport_place_giving

        template_name_id = DocumentSet.objects.get(pk=pk).template_name_id
        # .objects.get(pk=pk).template_name_id
        path_to_template = Sample.objects.get(pk=template_name_id).path_to_template.name

        file_name_from_path = os.path.basename(path_to_template)

        file_name_without_extension = str(os.path.splitext(file_name_from_path)[0])
        # file_name_without_extension = 'document_consent'

        # file_path = f'/Users/keito/Programming/Python/train/diploma/media/upload_sample/2024-05-07/{file_name_from_path}'
        # file_path = '/Users/keito/Programming/Python/train/diploma/media/upload_sample/2024-05-07/document_consent.docx'
        file_path = f'/Users/keito/Programming/Python/train/diploma/media/{path_to_template}'

        file_path_short = file_name_without_extension
        if file_path.endswith('.docx'):
            from datetime import datetime as dt
            doc = DocxTemplate(file_path)
            document2 = DocxTemplate(file_path)
            context = {
                'program_name': document_consent_program_name,
                'application_number': document_consent_application_number,
                'applicant_name': document_consent_applicant_name,
                'applicant_surname': document_consent_applicant_surname,
                'applicant_patronomic': document_consent_applicant_patronomic,
                'applicant_date_of_birth': dt.strftime(document_consent_applicant_date_of_birth, '%d.%m.%Y'),
                'birth_day': document_consent_applicant_date_of_birth.day,
                'birth_month': document_consent_applicant_date_of_birth.month if document_consent_applicant_date_of_birth.month >= 10 else str(
                    '0' + f'{document_consent_applicant_date_of_birth.month}'),
                'birth_year': document_consent_applicant_date_of_birth.year,
                'address_index': document_consent_address_index,
                'address_country': document_consent_address_country,
                'country_short': 'РФ',
                'address_city': document_consent_address_city,
                'address_street': document_consent_address_street,
                'address_building_number': document_consent_address_building_number,
                'address_house_flat_number': document_consent_address_house_flat_number,
                'applicant_phone_number': document_consent_applicant_phone_number,
                'current_date': dt.strftime(document_consent_current_date, '%d.%m.%Y'),
                'applicant_name_short': document_consent_applicant_name[0],
                'applicant_patronomic_short': document_consent_applicant_patronomic[0],
                'passport_seria': document_consent_passport_seria,
                'passport_number': document_consent_passport_number,
                'passport_date_of_issue': dt.strftime(document_consent_passport_date_of_issue, '%d.%m.%Y'),
                'passport_place_giving': document_consent_passport_place_giving,
            }
            import aspose.words as aw
            import random

            doc.render(context)
            document2.render(context)
            # file_name = f'/Users/keito/Downloads/{file_path_short}_downloaded.docx'
            file_name = f'{file_path_short}_downloaded.docx'
            # file_name_2 = f'{file_path_short}_downloaded.docx'
            # # document = aw.Document(file_name)
            # document2 = aw.Document(file_name_2)

            doc.save(file_name)
            # document2.save(file_name_2)

        response = FileResponse(open(file_name, 'rb'))
        response['Content-Disposition'] = f'attachment; filename="{file_name}"'

        # saved_document_word = DocumentConsent.objects.get(pk=pk)
        #
        # # saved_document_word.path_to_template = f'download_document/consent/2024-05-30/{file_name_2}'
        # saved_document_word.path_to_template = f'download_document/consent/2024-05-30/{file_name_2}'
        # saved_document_word.save()

        # response = FileResponse(open(file_name_2, 'rb'))
        # response['Content-Disposition'] = f'attachment; filename="{file_name_2}"'
        return response


class DocumentSetsDownloadPDF(View):
    def get(self, request, pk=1, *args, **kwargs):
        document_consent_program_name = DocumentSet.objects.get(pk=pk).document_consent.program_name
        # DocumentConsent.objects.get(pk=pk).program_name
        document_consent_application_number = DocumentSet.objects.get(pk=pk).document_consent.application_number
        document_consent_applicant_name = DocumentSet.objects.get(pk=pk).document_consent.applicant_name
        document_consent_applicant_surname = DocumentSet.objects.get(pk=pk).document_consent.applicant_surname
        document_consent_applicant_patronomic = DocumentSet.objects.get(pk=pk).document_consent.applicant_patronomic
        document_consent_applicant_date_of_birth = DocumentSet.objects.get(
            pk=pk).document_consent.applicant_date_of_birth
        document_consent_address_index = DocumentSet.objects.get(pk=pk).document_consent.address_index
        document_consent_address_country = DocumentSet.objects.get(pk=pk).document_consent.address_country
        document_consent_address_city = DocumentSet.objects.get(pk=pk).document_consent.address_city
        document_consent_address_street = DocumentSet.objects.get(pk=pk).document_consent.address_street
        document_consent_address_building_number = DocumentSet.objects.get(
            pk=pk).document_consent.address_building_number
        document_consent_address_house_flat_number = DocumentSet.objects.get(
            pk=pk).document_consent.address_house_flat_number
        document_consent_applicant_phone_number = DocumentSet.objects.get(pk=pk).document_consent.applicant_phone_number
        document_consent_current_date = DocumentSet.objects.get(pk=pk).document_consent.current_date
        document_consent_passport_seria = DocumentSet.objects.get(pk=pk).document_consent.passport_seria
        document_consent_passport_number = DocumentSet.objects.get(pk=pk).document_consent.passport_number
        document_consent_passport_date_of_issue = DocumentSet.objects.get(pk=pk).document_consent.passport_date_of_issue
        document_consent_passport_place_giving = DocumentSet.objects.get(pk=pk).document_consent.passport_place_giving

        template_name_id = DocumentSet.objects.get(pk=pk).template_name_id
        # .objects.get(pk=pk).template_name_id
        path_to_template = Sample.objects.get(pk=template_name_id).path_to_template.name

        file_name_from_path = os.path.basename(path_to_template)

        file_name_without_extension = str(os.path.splitext(file_name_from_path)[0])
        # file_name_without_extension = 'document_consent'

        # file_path = f'/Users/keito/Programming/Python/train/diploma/media/upload_sample/2024-05-07/{file_name_from_path}'
        # file_path = '/Users/keito/Programming/Python/train/diploma/media/upload_sample/2024-05-07/document_consent.docx'
        file_path = f'/Users/keito/Programming/Python/train/diploma/media/{path_to_template}'

        file_path_short = file_name_without_extension
        if file_path.endswith('.docx'):
            from datetime import datetime as dt
            doc = DocxTemplate(file_path)
            context = {
                'program_name': document_consent_program_name,
                'application_number': document_consent_application_number,
                'applicant_name': document_consent_applicant_name,
                'applicant_surname': document_consent_applicant_surname,
                'applicant_patronomic': document_consent_applicant_patronomic,
                'applicant_date_of_birth': dt.strftime(document_consent_applicant_date_of_birth, '%d.%m.%Y'),
                'birth_day': document_consent_applicant_date_of_birth.day,
                'birth_month': document_consent_applicant_date_of_birth.month if document_consent_applicant_date_of_birth.month >= 10 else str(
                    '0' + f'{document_consent_applicant_date_of_birth.month}'),
                'birth_year': document_consent_applicant_date_of_birth.year,
                'address_index': document_consent_address_index,
                'address_country': document_consent_address_country,
                'country_short': 'РФ',
                'address_city': document_consent_address_city,
                'address_street': document_consent_address_street,
                'address_building_number': document_consent_address_building_number,
                'address_house_flat_number': document_consent_address_house_flat_number,
                'applicant_phone_number': document_consent_applicant_phone_number,
                'current_date': dt.strftime(document_consent_current_date, '%d.%m.%Y'),
                'applicant_name_short': document_consent_applicant_name[0],
                'applicant_patronomic_short': document_consent_applicant_patronomic[0],
                'passport_seria': document_consent_passport_seria,
                'passport_number': document_consent_passport_number,
                'passport_date_of_issue': dt.strftime(document_consent_passport_date_of_issue, '%d.%m.%Y'),
                'passport_place_giving': document_consent_passport_place_giving,
            }

            doc.render(context)
            file_name = f'{file_path_short}_downloaded.docx'
            import aspose.words as aw
            file_name_pdf = f'{file_path_short}_downloaded.pdf'
            document = aw.Document(file_name)
            # doc.save(file_name)
            document.save(file_name_pdf)

        response = FileResponse(open(file_name_pdf, 'rb'))
        response['Content-Disposition'] = f'attachment; filename="{file_name_pdf}"'
        return response


class SampleCreateView(CreateView):
    model = Sample
    form_class = SampleForm
    template_name = 'sample_new.html'

    # fields = ['title', 'path_to_template']
    success_url = reverse_lazy('samples_all')


class SampleUpdateView(UpdateView):
    model = Sample
    form_class = SampleForm
    template_name = 'sample_edit.html'
    # fields = ['title', 'path_to_template']


class SampleDeleteView(DeleteView):
    model = Sample
    template_name = 'sample_delete.html'
    success_url = reverse_lazy('samples_all')


class ReviewProblemSampleCreateView(CreateView):
    model = ReviewProblemSample
    form_class = ReviewProblemSampleForm
    template_name = 'ReviewProblemSample/review_problem_sample_new.html'

    # fields = ['title', 'path_to_template']
    success_url = reverse_lazy('review_problem_sample_success')


class ReviewProblemDocumentConsentCreateView(CreateView):
    model = ReviewProblemDocumentConsent
    form_class = ReviewProblemDocumentConsentForm

    template_name = 'ReviewProblemDocument/review_problem_document_consent_new.html'

    # fields = ['title', 'path_to_template']
    success_url = reverse_lazy('review_problem_document_consent_success')


class SignDocumentCreateView(CreateView):
    model = SignDocument
    form_class = SignDocumentForm
    template_name = 'document_sign/sign_dc_new.html'

    # fields = ['title', 'path_to_template']
    success_url = reverse_lazy('head_sign_dc_all')


class SignDocumentListView(ListView):
    model = SignDocument
    template_name = 'document_sign/sign_dc_all.html'


class SignDocumentDetailView(DetailView):
    model = SignDocument
    template_name = 'document_sign/sign_dc_detail.html'


class SignDocumentDeleteView(DeleteView):
    model = SignDocument
    template_name = 'document_sign/sign_dc_delete.html'
    success_url = reverse_lazy('sign_dc_all')


# class SignDocumentUpdateView(UpdateView):
#     model = Sample
#     form_class = SignDocumentForm
#     template_name = 'document_sign/sign_dc_edit.html'
#     # fields = ['title', 'path_to_template']


class HeadSignDocumentCreateView(CreateView):
    model = SignDocument
    form_class = SignDocumentForm
    template_name = 'document_sign/head_sign_dc/head_sign_dc_new.html'

    # fields = ['title', 'path_to_template']
    success_url = reverse_lazy('head_sign_dc_all')


class HeadSignDocumentListView(ListView):
    model = SignDocument
    template_name = 'document_sign/head_sign_dc/head_sign_dc_all.html'


class HeadSignDocumentDetailView(DetailView):
    model = SignDocument
    template_name = 'document_sign/head_sign_dc/head_sign_dc_detail.html'


class HeadSignDocumentDeleteView(DeleteView):
    model = SignDocument
    template_name = 'document_sign/head_sign_dc/head_sign_dc_delete.html'
    success_url = reverse_lazy('sign_dc_all')


class HeadSignDocumentUpdateView(UpdateView):
    model = SignDocument
    form_class = SignDocumentForm
    # fields = ['status', ]
    template_name = 'document_sign/head_sign_dc/head_sign_dc_edit.html'
    success_url = reverse_lazy('head_sign_dc_all')

    # fields = ['title', 'path_to_template']
